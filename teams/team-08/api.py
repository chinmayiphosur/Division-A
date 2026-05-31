from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import json
import csv
import os
import time
from dotenv import load_dotenv
load_dotenv()
from datetime import datetime
import io
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from groq import Groq
from logger import get_logger, log_incident, log_api_call

# Use environment variable for the API key to prevent exposing it to GitHub
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
client = Groq(api_key=GROQ_API_KEY)

# ============================================================
# CONFIGURATION & LOGGING INTERFACE
# ============================================================

logger = get_logger("api_server")
app = FastAPI(title="Breach Analyser API")

# Request Telemetry Middleware
@app.middleware("http")
async def log_requests(request: Request, call_next):
    start_time = time.time()
    response = await call_next(request)
    process_time = (time.time() - start_time) * 1000
    
    # Do not log the live log viewer endpoint to avoid endless loops of logs!
    if "/api/logs" not in request.url.path:
        log_api_call(logger, request.method, request.url.path, response.status_code, process_time)
        
    return response

# Enable CORS for the React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

MODEL_NAME = "llama-3.3-70b-versatile"

DATA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")

# ============================================================
# DATA LOADERS & DATASET TELEMETRY
# ============================================================

def load_pii_config():
    logger.info("Initializing PII field configuration database...")
    start_time = time.time()
    try:
        with open(os.path.join(DATA_DIR, "pii_field_config.json"), "r") as f:
            data = json.load(f)
        logger.info(f"Loaded PII configuration ({len(data)} fields) successfully in {(time.time() - start_time)*1000:.2f}ms")
        return data
    except Exception as e:
        logger.error(f"Failed to load PII field configuration: {str(e)}", exc_info=True)
        raise e

def load_compliance_rules():
    logger.info("Initializing compliance framework rules...")
    start_time = time.time()
    try:
        with open(os.path.join(DATA_DIR, "compliance_rules.json"), "r") as f:
            data = json.load(f)
        logger.info(f"Loaded compliance rules (GDPR & DPDPA) successfully in {(time.time() - start_time)*1000:.2f}ms")
        return data
    except Exception as e:
        logger.error(f"Failed to load compliance rules: {str(e)}", exc_info=True)
        raise e

def load_hibp_stats():
    logger.info("Initializing HaveIBeenPwned stats aggregation...")
    start_time = time.time()
    stats = {}
    try:
        with open(os.path.join(DATA_DIR, "hibp_breaches_full.json"), "r", encoding="utf-8") as f:
            breaches = json.load(f)
    except FileNotFoundError:
        logger.warning("hibp_breaches_full.json not found! HaveIBeenPwned statistics will be disabled.")
        return {}
    except Exception as e:
        logger.error(f"Failed to load HIBP breaches: {str(e)}", exc_info=True)
        return {}
        
    FIELD_MAPPING = {
        "email": ["Email addresses"],
        "password": ["Passwords", "Password hints"],
        "username": ["Usernames"],
        "ip_address": ["IP addresses"],
        "medical_records": ["Health insurance information", "Medical records"],
        "credit_card": ["Credit cards"],
        "ssn_aadhaar": ["Social security numbers", "Government issued IDs", "Nationalities"],
        "passport": ["Passport numbers"],
        "bank_account": ["Bank account numbers", "Financial transactions"],
        "dob": ["Dates of birth"],
        "phone": ["Phone numbers"],
        "address": ["Physical addresses"],
        "biometric": ["Biometric data"]
    }
    
    for field, hibp_classes in FIELD_MAPPING.items():
        field_breaches = [b for b in breaches if any(c in b.get("DataClasses", []) for c in hibp_classes)]
        if field_breaches:
            biggest = max(field_breaches, key=lambda x: x.get("PwnCount", 0))
            total_records = sum(b.get("PwnCount", 0) for b in field_breaches)
            recent_year = max(int(b.get("BreachDate", "2000").split("-")[0]) for b in field_breaches if b.get("BreachDate"))
            stats[field] = {
                "breach_count": len(field_breaches),
                "biggest_breach": biggest.get("Name", "Unknown"),
                "biggest_records": biggest.get("PwnCount", 0),
                "most_recent": str(recent_year),
                "avg_records": total_records // len(field_breaches),
                "source": "HaveIBeenPwned API"
            }
    logger.info(f"Aggregated HaveIBeenPwned database ({len(breaches)} entries, {len(stats)} fields mapped) in {(time.time() - start_time)*1000:.2f}ms")
    return stats

def load_gdpr_fines():
    logger.info("Loading GDPR enforcement tracker fines...")
    start_time = time.time()
    fines = {}
    try:
        with open(os.path.join(DATA_DIR, "gdpr_enforcement_tracker.csv"), "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                fines[row["field"]] = {
                    "example_org": row["example_org"],
                    "country": row["country"],
                    "fine": f"€{int(float(row['fine_eur_equivalent'])):,}",
                    "fine_raw": int(float(row['fine_eur_equivalent'])),
                    "year": row["year"],
                    "articles": row["articles_violated"],
                    "authority": row["authority"],
                    "source": row["source"]
                }
        logger.info(f"Loaded GDPR fines tracker ({len(fines)} precedent categories) in {(time.time() - start_time)*1000:.2f}ms")
    except Exception as e:
        logger.error(f"Failed to load GDPR fines database: {str(e)}", exc_info=True)
    return fines

def load_sector_benchmarks():
    logger.info("Loading industry sector cost benchmarks...")
    start_time = time.time()
    benchmarks = {}
    try:
        with open(os.path.join(DATA_DIR, "sector_benchmarks.csv"), "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                benchmarks[row["sector"]] = {
                    "avg_breach_cost_usd": f"{int(row['avg_breach_cost_usd']) / 1_000_000:.1f}M",
                    "avg_records_lost": row["avg_records_lost"],
                    "most_common_attack": row["most_common_attack"],
                    "avg_detection_days": int(row["avg_detection_days"]),
                    "avg_containment_days": int(row["avg_containment_days"]),
                    "per_record_cost": float(row["per_record_cost_usd"]),
                    "source": row["source"]
                }
        logger.info(f"Loaded industry sector benchmarks ({len(benchmarks)} sectors mapped) in {(time.time() - start_time)*1000:.2f}ms")
    except Exception as e:
        logger.error(f"Failed to load sector benchmarks: {str(e)}", exc_info=True)
    return benchmarks

def load_privacy_rights_breaches():
    logger.info("Loading Privacy Rights Clearinghouse archive...")
    start_time = time.time()
    breaches = []
    try:
        with open(os.path.join(DATA_DIR, "privacy_rights_clearinghouse.csv"), "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                breaches.append({
                    "name": row["breach_name"],
                    "org": row["organization"],
                    "sector": row["sector"],
                    "year": int(row["year"]),
                    "records": int(row["records_lost"]),
                    "fields": [x.strip() for x in row["data_types_exposed"].split(",")],
                    "country": row["state_country"],
                    "source": row["source"]
                })
        logger.info(f"Loaded PRC database archive ({len(breaches)} historical breach logs) in {(time.time() - start_time)*1000:.2f}ms")
    except Exception as e:
        logger.error(f"Failed to load PRC historical archive: {str(e)}", exc_info=True)
    return breaches

# Load globals
PII_CONFIG = load_pii_config()
COMPLIANCE_RULES = load_compliance_rules()
HIBP_STATS = load_hibp_stats()
GDPR_FINES = load_gdpr_fines()
SECTOR_BENCHMARKS = load_sector_benchmarks()
PRC_BREACHES = load_privacy_rights_breaches()

FIELD_WEIGHTS = {k: v["weight"] for k, v in PII_CONFIG.items()}
PII_SENSITIVITY = {k: v["sensitivity"] for k, v in PII_CONFIG.items()}
FIELD_EXPLANATIONS = {k: v["explanation"] for k, v in PII_CONFIG.items()}
SECTOR_MULTIPLIER = COMPLIANCE_RULES["sector_multipliers"]
GDPR_FIELD_TRIGGERS = COMPLIANCE_RULES["gdpr"]["field_triggers"]
DPDPA_FIELD_TRIGGERS = COMPLIANCE_RULES["dpdpa"]["field_triggers"]

GDPR_RULES = "\n".join([f"{k}: {v}" for k, v in COMPLIANCE_RULES["gdpr"]["full_text"].items()])
DPDPA_RULES = "\n".join([f"{k}: {v}" for k, v in COMPLIANCE_RULES["dpdpa"]["full_text"].items()])

# ============================================================
# API MODELS
# ============================================================

class AnalyzeRequest(BaseModel):
    org_name: str
    sector: str
    breach_date: str
    num_users: int
    jurisdiction: str
    fields: List[str]

# ============================================================
# ENDPOINTS
# ============================================================

@app.get("/api/config")
def get_config():
    """Return available options for the frontend UI"""
    return {
        "fields": [
            {"id": k, "label": k.replace("_", " ").title(), "sensitivity": PII_SENSITIVITY.get(k, "Medium")}
            for k in FIELD_WEIGHTS.keys()
        ],
        "sectors": list(SECTOR_MULTIPLIER.keys()),
        "jurisdictions": ["India", "EU", "India + EU"]
    }

@app.get("/api/stats")
def get_stats():
    """Return stats for the historical dashboard"""
    return {
        "hibp": HIBP_STATS,
        "gdpr": GDPR_FINES,
        "benchmarks": SECTOR_BENCHMARKS
    }

@app.post("/api/analyze")
def analyze_breach(req: AnalyzeRequest):
    logger.info(f"Received breach analysis request for organization: '{req.org_name}' [Sector: {req.sector}, Jurisdiction: {req.jurisdiction}, Users: {req.num_users:,}]")
    start_time = time.time()
    try:
        # Risk calculation
        weights = sorted([FIELD_WEIGHTS.get(f, 3) for f in req.fields], reverse=True)
        if len(weights) == 0:
            base = 0
        elif len(weights) == 1:
            base = weights[0]
        elif len(weights) == 2:
            base = weights[0] * 0.7 + weights[1] * 0.3
        else:
            base = weights[0] * 0.5 + weights[1] * 0.3 + weights[2] * 0.2

        volume_bonus = min(2, req.num_users / 500000)
        sector_mult = SECTOR_MULTIPLIER.get(req.sector, 1.0)
        score = min(10, (base + volume_bonus) * sector_mult)
        score = round(score, 1)

        if score < 3: severity = "Low"
        elif score < 6: severity = "Medium"
        elif score < 8: severity = "High"
        else: severity = "Critical"

        logger.info(f"Risk engine calculated score: {score}/10 (Severity: {severity}) for '{req.org_name}'")

        # Compliance
        gdpr_articles = set()
        dpdpa_sections = set()
        for field in req.fields:
            gdpr_articles.update(GDPR_FIELD_TRIGGERS.get(field, []))
            dpdpa_sections.update(DPDPA_FIELD_TRIGGERS.get(field, []))
        gdpr_articles = sorted(list(gdpr_articles))
        dpdpa_sections = sorted(list(dpdpa_sections))

        # Explanations
        explanations = []
        for field in req.fields:
            explanations.append({
                "field": field.replace("_", " ").title(),
                "raw_field": field,
                "sensitivity": PII_SENSITIVITY.get(field, "Medium"),
                "risk": FIELD_EXPLANATIONS.get(field, "Potential privacy risk"),
                "category": PII_CONFIG.get(field, {}).get("category", "Other")
            })

        # Financial Impact
        bench = SECTOR_BENCHMARKS.get(req.sector, {})
        base_cost = bench.get("per_record_cost", 4.0)
        FIELD_COST_MULTIPLIER = {
            "medical_records": 2.5, "credit_card": 2.0, "biometric": 2.2,
            "ssn_aadhaar": 1.8, "bank_account": 1.9, "password": 1.3,
            "passport": 1.6, "email": 1.0, "phone": 1.1, "dob": 1.2,
            "address": 1.1, "username": 0.8, "ip_address": 0.7
        }
        max_mult = max([FIELD_COST_MULTIPLIER.get(f, 1.0) for f in req.fields] + [1.0])
        estimated = req.num_users * base_cost * max_mult

        financial_impact = {
            "min": f"${estimated * 0.7:,.0f}",
            "max": f"${estimated * 1.3:,.0f}",
            "likely": f"${estimated:,.0f}"
        }

        # Similar Breaches
        scored = []
        field_set = set(req.fields)
        for breach in PRC_BREACHES:
            breach_fields = set(breach["fields"])
            overlap = len(field_set & breach_fields)
            union = len(field_set | breach_fields)
            field_similarity = overlap / union if union > 0 else 0
            sector_bonus = 0.3 if breach["sector"] == req.sector else 0
            scale_ratio = min(breach["records"], req.num_users) / max(breach["records"], req.num_users) if breach["records"] > 0 and req.num_users > 0 else 0
            total_score = field_similarity * 0.5 + sector_bonus + scale_ratio * 0.2
            if overlap > 0:
                scored.append({
                    **breach,
                    "similarity_score": round(total_score * 100, 1),
                    "overlap_fields": sorted(list(field_set & breach_fields)),
                    "overlap_count": overlap
                })
        scored.sort(key=lambda x: x["similarity_score"], reverse=True)
        similar_breaches = scored[:5]

        # LLM Analysis
        fields_str = ", ".join([f.replace("_", " ").title() for f in req.fields])
        gdpr_str = ", ".join(gdpr_articles)
        dpdpa_str = ", ".join(dpdpa_sections)

        hibp_lines = []
        for field in req.fields:
            if field in HIBP_STATS:
                s = HIBP_STATS[field]
                hibp_lines.append(
                    f"- {field.replace('_', ' ').title()}: appeared in {s['breach_count']} known breaches, biggest was {s['biggest_breach']} ({s['biggest_records']:,} records)"
                )
        hibp_context = "\n".join(hibp_lines)

        similar_lines = []
        for b in similar_breaches[:3]:
            similar_lines.append(f"- {b['name']} ({b['year']}): {b['records']:,} records, sector: {b['sector']}, fields: {', '.join(b['overlap_fields'])}")
        similar_context = "\n".join(similar_lines) if similar_lines else "No similar breaches found."

        prompt = f"""
You are a senior cybersecurity compliance analyst.
A data breach has occurred. Here are the facts:

Organisation: {req.org_name}
Sector: {req.sector}
Breach Date: {req.breach_date}
Affected Users: {req.num_users:,}
Jurisdiction: {req.jurisdiction}
Breached Data Fields: {fields_str}
Risk Score: {score}/10 (Severity: {severity})
GDPR Articles Triggered: {gdpr_str}
DPDPA Sections Triggered: {dpdpa_str}

Historical Breach Intelligence (HaveIBeenPwned database):
{hibp_context}

Similar Past Breaches (Privacy Rights Clearinghouse database):
{similar_context}

GDPR Reference:
{GDPR_RULES}

DPDPA 2023 Reference:
{DPDPA_RULES}

Your task: Return a JSON object with EXACTLY these keys:
{{
    "executive_summary": "3-4 sentence summary of the breach impact, referencing similar past breaches",
    "immediate_actions": ["action1", "action2", "action3", "action4"],
    "gdpr_analysis": "2-3 sentences explaining GDPR obligations for this specific breach",
    "dpdpa_analysis": "2-3 sentences explaining DPDPA obligations for this specific breach",
    "risk_justification": "2-3 sentences explaining why this risk score was assigned, citing historical data",
    "disclosure_letter": "Full professional breach notification letter addressed to affected individuals"
}}
Return ONLY valid JSON. No markdown. No explanation outside JSON.
"""

        logger.info(f"Submitting context-augmented breach payload to Groq Llama model '{MODEL_NAME}'...")
        llm_start_time = time.time()
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=4096
        )
        text = response.choices[0].message.content.strip()
        llm_duration_ms = (time.time() - llm_start_time) * 1000
        logger.info(f"Groq intelligence analysis generated successfully in {llm_duration_ms:.2f}ms")

        if text.startswith("```"):
            text = text.split("```")[1]
            if text.startswith("json"):
                text = text[4:]
        llm_output = json.loads(text.strip())

        gdpr_fines_triggered = {f: GDPR_FINES[f] for f in req.fields if f in GDPR_FINES}
        hibp_stats_triggered = {f: HIBP_STATS[f] for f in req.fields if f in HIBP_STATS}

        total_duration_ms = (time.time() - start_time) * 1000
        log_incident(logger, req.org_name, req.sector, req.num_users, score, severity, total_duration_ms)

        return {
            "risk": {
                "score": score,
                "severity": severity,
                "financial_impact": financial_impact,
                "benchmarks": bench
            },
            "compliance": {
                "gdpr_articles": gdpr_articles,
                "dpdpa_sections": dpdpa_sections,
                "explanations": explanations,
                "fines": gdpr_fines_triggered
            },
            "intelligence": {
                "hibp_stats": hibp_stats_triggered
            },
            "similar_breaches": similar_breaches,
            "llm_analysis": llm_output
        }
    except Exception as e:
        logger.error(f"Failed to analyze data breach for '{req.org_name}': {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Breach analysis processing error: {str(e)}")

class ReportRequest(BaseModel):
    org_name: str
    breach_date: str
    num_users: int
    sector: str
    fields: List[str]
    risk_score: float
    severity: str
    gdpr_articles: List[str]
    dpdpa_sections: List[str]
    financial_impact: Dict[str, str]
    similar_breaches: List[Dict[str, Any]]
    llm_analysis: Dict[str, Any]

@app.post("/api/report")
def generate_report(req: ReportRequest):
    logger.info(f"Received request to compile incident report document for '{req.org_name}'")
    start_time = time.time()
    try:
        doc = Document()
        
        title = doc.add_heading("DATA BREACH INCIDENT REPORT", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}")
        doc.add_paragraph(f"Organisation: {req.org_name}")
        doc.add_paragraph(f"Sector: {req.sector.title()}")
        doc.add_paragraph(f"Breach Date: {req.breach_date}")
        doc.add_paragraph(f"Affected Users: {req.num_users:,}")
        doc.add_paragraph("")

        doc.add_heading("RISK ASSESSMENT", level=1)
        doc.add_paragraph(f"Risk Score: {req.risk_score}/10")
        doc.add_paragraph(f"Severity: {req.severity}")

        doc.add_heading("ESTIMATED FINANCIAL IMPACT", level=1)
        doc.add_paragraph(f"Minimum Estimate: {req.financial_impact['min']}")
        doc.add_paragraph(f"Most Likely: {req.financial_impact['likely']}")
        doc.add_paragraph(f"Maximum Estimate: {req.financial_impact['max']}")

        doc.add_heading("BREACHED DATA FIELDS", level=1)
        for field in req.fields:
            sens = PII_SENSITIVITY.get(field, "Medium")
            exp = FIELD_EXPLANATIONS.get(field, "Potential privacy risk")
            cat = PII_CONFIG.get(field, {}).get("category", "Other")
            doc.add_paragraph(f"• {field.replace('_', ' ').title()} [{cat}] — Sensitivity: {sens} — {exp}", style="List Bullet")

        doc.add_heading("HISTORICAL BREACH INTELLIGENCE", level=1)
        for field in req.fields:
            if field in HIBP_STATS:
                s = HIBP_STATS[field]
                doc.add_paragraph(f"• {field.replace('_', ' ').title()}: {s['breach_count']} known breaches, biggest: {s['biggest_breach']} ({s['biggest_records']:,} records)", style="List Bullet")

        doc.add_heading("SIMILAR PAST BREACHES", level=1)
        for b in req.similar_breaches:
            doc.add_paragraph(f"• {b['name']} ({b['org']}, {b['year']}): {b['records']:,} records — Similarity: {b['similarity_score']}%", style="List Bullet")

        doc.add_heading("REGULATORY OBLIGATIONS", level=1)
        doc.add_heading("GDPR", level=2)
        doc.add_paragraph(req.llm_analysis['gdpr_analysis'])
        for article in req.gdpr_articles:
            doc.add_paragraph(f"• {article}", style="List Bullet")
            
        doc.add_heading("GDPR Fine Precedents", level=2)
        for field in req.fields:
            if field in GDPR_FINES:
                ref = GDPR_FINES[field]
                doc.add_paragraph(f"• {field.replace('_', ' ').title()}: {ref['example_org']} fined {ref['fine']} ({ref['year']}) — {ref['articles']}", style="List Bullet")

        doc.add_heading("DPDPA 2023", level=2)
        doc.add_paragraph(req.llm_analysis['dpdpa_analysis'])
        for section in req.dpdpa_sections:
            doc.add_paragraph(f"• {section}", style="List Bullet")

        doc.add_heading("IMMEDIATE ACTIONS REQUIRED", level=1)
        for action in req.llm_analysis['immediate_actions']:
            doc.add_paragraph(f"• {action}", style="List Bullet")

        doc.add_heading("BREACH NOTIFICATION LETTER", level=1)
        doc.add_paragraph(req.llm_analysis['disclosure_letter'])

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        duration_ms = (time.time() - start_time) * 1000
        logger.info(f"Compiled and packaged DOCX report document for '{req.org_name}' successfully in {duration_ms:.2f}ms")
        
        return StreamingResponse(
            buffer, 
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=BreachReport_{req.org_name}.docx"}
        )
    except Exception as e:
        logger.error(f"Failed to generate report document for '{req.org_name}': {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")

@app.get("/api/logs")
def get_logs(limit: int = 100, level: Optional[str] = None):
    """
    Audit log retrieval endpoint. Parses the persistent rotating log file
    into structured JSON telemetry objects for system observation.
    """
    from logger import LOG_FILE_PATH
    if not os.path.exists(LOG_FILE_PATH):
        return []
        
    parsed_logs = []
    try:
        with open(LOG_FILE_PATH, "r", encoding="utf-8") as f:
            lines = f.readlines()
            
        for line in reversed(lines):
            line = line.strip()
            if not line:
                continue
                
            parts = line.split(" | ")
            if len(parts) >= 4:
                timestamp = parts[0].strip()
                lvl = parts[1].strip()
                module = parts[2].strip()
                
                message_part = " | ".join(parts[3:])
                metadata = None
                
                # Check for metadata signature
                if " | {" in message_part:
                    msg_text, meta_text = message_part.split(" | {", 1)
                    meta_text = "{" + meta_text
                    try:
                        metadata = json.loads(meta_text)
                        message_part = msg_text
                    except json.JSONDecodeError:
                        pass
                
                if level and lvl != level.upper():
                    continue
                    
                parsed_logs.append({
                    "timestamp": timestamp,
                    "level": lvl,
                    "module": module,
                    "message": message_part,
                    "metadata": metadata
                })
                
                if len(parsed_logs) >= limit:
                    break
    except Exception as e:
        # Fallback console reporting
        print(f"Log parsing error: {str(e)}")
        return [{
            "timestamp": str(datetime.now()),
            "level": "ERROR",
            "module": "api_server",
            "message": f"Log subsystem extraction failure: {str(e)}",
            "metadata": None
        }]
        
    return parsed_logs

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("api:app", host="127.0.0.1", port=8000, reload=True)
