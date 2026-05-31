import streamlit as st
from groq import Groq
import json
import csv
import os
import time
from dotenv import load_dotenv
load_dotenv()
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
import io
import plotly.graph_objects as go
from logger import get_logger, log_incident

# ============================================================
# CONFIGURATION & LOGGING INTERFACE
# ============================================================

logger = get_logger("streamlit_app")
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
client = Groq(api_key=GROQ_API_KEY)
MODEL_NAME = "llama-3.3-70b-versatile"

DATA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")

# ============================================================
# DATA LOADERS — Loading from external datasets
# ============================================================

def load_pii_config():
    """Load PII field configuration from data/pii_field_config.json"""
    logger.info("Streamlit: Loading PII field configuration...")
    t0 = time.time()
    try:
        with open(os.path.join(DATA_DIR, "pii_field_config.json"), "r") as f:
            data = json.load(f)
        logger.info(f"Streamlit: Successfully loaded PII config ({len(data)} fields) in {(time.time() - t0)*1000:.2f}ms")
        return data
    except Exception as e:
        logger.error(f"Streamlit: Failed to load PII configuration: {str(e)}", exc_info=True)
        raise e

def load_compliance_rules():
    """Load GDPR/DPDPA rules from data/compliance_rules.json"""
    logger.info("Streamlit: Loading compliance framework rules...")
    t0 = time.time()
    try:
        with open(os.path.join(DATA_DIR, "compliance_rules.json"), "r") as f:
            data = json.load(f)
        logger.info(f"Streamlit: Successfully loaded compliance rules in {(time.time() - t0)*1000:.2f}ms")
        return data
    except Exception as e:
        logger.error(f"Streamlit: Failed to load compliance rules: {str(e)}", exc_info=True)
        raise e

def load_hibp_stats():
    """Dynamically aggregate stats from the full HIBP JSON dataset"""
    logger.info("Streamlit: Initializing HaveIBeenPwned stats aggregator...")
    t0 = time.time()
    stats = {}
    
    try:
        with open(os.path.join(DATA_DIR, "hibp_breaches_full.json"), "r", encoding="utf-8") as f:
            breaches = json.load(f)
    except FileNotFoundError:
        logger.warning("Streamlit: hibp_breaches_full.json not found! HIBP stats will be disabled.")
        return {}
    except Exception as e:
        logger.error(f"Streamlit: Failed to load HIBP breaches: {str(e)}", exc_info=True)
        return {}
        
    FIELD_MAPPING = {
        "email": ["Email addresses"],
        "password": ["Passwords", "Password hints"],
        "username": ["Usernames"],
        "ip_address": ["IP addresses"],
        "medical_records": ["Health insurance information", "Medical records"],
        "credit_card": ["Credit cards"],
        "ssn_aadhaar": ["Social security numbers", "Government issued IDs", "Nationalities"],
        "passport": ["Passport numbers"],
        "bank_account": ["Bank account numbers", "Financial transactions"],
        "dob": ["Dates of birth"],
        "phone": ["Phone numbers"],
        "address": ["Physical addresses"],
        "biometric": ["Biometric data"]
    }
    
    for field, hibp_classes in FIELD_MAPPING.items():
        field_breaches = [b for b in breaches if any(c in b.get("DataClasses", []) for c in hibp_classes)]
        
        if field_breaches:
            biggest = max(field_breaches, key=lambda x: x.get("PwnCount", 0))
            total_records = sum(b.get("PwnCount", 0) for b in field_breaches)
            recent_year = max(int(b.get("BreachDate", "2000").split("-")[0]) for b in field_breaches if b.get("BreachDate"))
            
            stats[field] = {
                "breach_count": len(field_breaches),
                "biggest_breach": biggest.get("Name", "Unknown"),
                "biggest_records": biggest.get("PwnCount", 0),
                "most_recent": str(recent_year),
                "avg_records": total_records // len(field_breaches),
                "source": "HaveIBeenPwned API (Full Dataset)"
            }
    logger.info(f"Streamlit: Successfully aggregated HIBP statistics ({len(breaches)} records, {len(stats)} fields mapped) in {(time.time() - t0)*1000:.2f}ms")
    return stats

def load_gdpr_fines():
    """Load GDPR enforcement tracker from data/gdpr_enforcement_tracker.csv"""
    logger.info("Streamlit: Loading GDPR fines precedents dataset...")
    t0 = time.time()
    fines = {}
    try:
        with open(os.path.join(DATA_DIR, "gdpr_enforcement_tracker.csv"), "r",
                  encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                fines[row["field"]] = {
                    "example_org": row["example_org"],
                    "country": row["country"],
                    "fine": f"€{int(float(row['fine_eur_equivalent'])):,}",
                    "fine_raw": int(float(row['fine_eur_equivalent'])),
                    "year": row["year"],
                    "articles": row["articles_violated"],
                    "authority": row["authority"],
                    "source": row["source"]
                }
        logger.info(f"Streamlit: Successfully loaded GDPR fines ({len(fines)} categories) in {(time.time() - t0)*1000:.2f}ms")
    except Exception as e:
        logger.error(f"Streamlit: Failed to load GDPR fines database: {str(e)}", exc_info=True)
    return fines

def load_sector_benchmarks():
    """Load sector benchmarks from data/sector_benchmarks.csv"""
    logger.info("Streamlit: Loading industry cost benchmarks...")
    t0 = time.time()
    benchmarks = {}
    try:
        with open(os.path.join(DATA_DIR, "sector_benchmarks.csv"), "r",
                  encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                benchmarks[row["sector"]] = {
                    "avg_breach_cost_usd": f"{int(row['avg_breach_cost_usd']) / 1_000_000:.1f}M",
                    "avg_records_lost": row["avg_records_lost"],
                    "most_common_attack": row["most_common_attack"],
                    "avg_detection_days": int(row["avg_detection_days"]),
                    "avg_containment_days": int(row["avg_containment_days"]),
                    "per_record_cost": float(row["per_record_cost_usd"]),
                    "source": row["source"]
                }
        logger.info(f"Streamlit: Successfully loaded sector cost benchmarks ({len(benchmarks)} sectors) in {(time.time() - t0)*1000:.2f}ms")
    except Exception as e:
        logger.error(f"Streamlit: Failed to load sector benchmarks: {str(e)}", exc_info=True)
    return benchmarks

def load_privacy_rights_breaches():
    """Load Privacy Rights Clearinghouse breach records from CSV"""
    logger.info("Streamlit: Loading Privacy Rights Clearinghouse historical archive...")
    t0 = time.time()
    breaches = []
    try:
        with open(os.path.join(DATA_DIR, "privacy_rights_clearinghouse.csv"), "r",
                  encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                breaches.append({
                    "name": row["breach_name"],
                    "org": row["organization"],
                    "sector": row["sector"],
                    "year": int(row["year"]),
                    "records": int(row["records_lost"]),
                    "fields": [x.strip() for x in row["data_types_exposed"].split(",")],
                    "country": row["state_country"],
                    "source": row["source"]
                })
        logger.info(f"Streamlit: Successfully loaded PRC historical archive ({len(breaches)} breach records) in {(time.time() - t0)*1000:.2f}ms")
    except Exception as e:
        logger.error(f"Streamlit: Failed to load PRC historical archive: {str(e)}", exc_info=True)
    return breaches

# ============================================================
# LOAD ALL DATASETS AT STARTUP
# ============================================================

@st.cache_data
def load_all_datasets():
    """Load and cache all datasets from the data/ directory"""
    pii_config = load_pii_config()
    compliance = load_compliance_rules()
    hibp_stats = load_hibp_stats()
    gdpr_fines = load_gdpr_fines()
    sector_bench = load_sector_benchmarks()
    prc_breaches = load_privacy_rights_breaches()
    return pii_config, compliance, hibp_stats, gdpr_fines, sector_bench, prc_breaches

(PII_CONFIG, COMPLIANCE_RULES, HIBP_STATS,
 GDPR_FINES, SECTOR_BENCHMARKS, PRC_BREACHES) = load_all_datasets()

# Build derived lookup dictionaries from loaded data
FIELD_WEIGHTS = {k: v["weight"] for k, v in PII_CONFIG.items()}
PII_SENSITIVITY = {k: v["sensitivity"] for k, v in PII_CONFIG.items()}
FIELD_EXPLANATIONS = {k: v["explanation"] for k, v in PII_CONFIG.items()}
SECTOR_MULTIPLIER = COMPLIANCE_RULES["sector_multipliers"]
GDPR_FIELD_TRIGGERS = COMPLIANCE_RULES["gdpr"]["field_triggers"]
DPDPA_FIELD_TRIGGERS = COMPLIANCE_RULES["dpdpa"]["field_triggers"]

# Build compliance text strings from loaded JSON
GDPR_RULES = "\n".join(
    [f"{k}: {v}" for k, v in COMPLIANCE_RULES["gdpr"]["full_text"].items()]
)
DPDPA_RULES = "\n".join(
    [f"{k}: {v}" for k, v in COMPLIANCE_RULES["dpdpa"]["full_text"].items()]
)

# ============================================================
# RISK ENGINE
# ============================================================

def calculate_risk(fields, sector, num_users):
    weights = sorted(
        [FIELD_WEIGHTS.get(f, 3) for f in fields],
        reverse=True
    )
    if len(weights) == 1:
        base = weights[0]
    elif len(weights) == 2:
        base = weights[0] * 0.7 + weights[1] * 0.3
    else:
        base = weights[0] * 0.5 + weights[1] * 0.3 + weights[2] * 0.2

    volume_bonus = min(2, num_users / 500000)
    sector_mult = SECTOR_MULTIPLIER.get(sector, 1.0)
    score = min(10, (base + volume_bonus) * sector_mult)
    score = round(score, 1)

    if score < 3:
        severity = "Low"
        color = "green"
    elif score < 6:
        severity = "Medium"
        color = "orange"
    elif score < 8:
        severity = "High"
        color = "red"
    else:
        severity = "Critical"
        color = "darkred"

    return score, severity, color


def get_compliance_triggers(fields):
    gdpr_articles = set()
    dpdpa_sections = set()
    for field in fields:
        gdpr_articles.update(GDPR_FIELD_TRIGGERS.get(field, []))
        dpdpa_sections.update(DPDPA_FIELD_TRIGGERS.get(field, []))
    return sorted(list(gdpr_articles)), sorted(list(dpdpa_sections))


def generate_explanation(fields):
    lines = []
    for field in fields:
        exp = FIELD_EXPLANATIONS.get(field, "Potential privacy risk")
        sens = PII_SENSITIVITY.get(field, "Medium")
        cat = PII_CONFIG.get(field, {}).get("category", "Other")
        lines.append({
            "field": field.replace("_", " ").title(),
            "raw_field": field,
            "sensitivity": sens,
            "risk": exp,
            "category": cat
        })
    return lines

# ============================================================
# SIMILAR BREACH FINDER (Privacy Rights Clearinghouse)
# ============================================================

def find_similar_breaches(fields, sector, num_users, top_n=5):
    """Find similar historical breaches from the PRC database"""
    scored = []
    field_set = set(fields)

    for breach in PRC_BREACHES:
        breach_fields = set(breach["fields"])
        # Jaccard similarity for field overlap
        overlap = len(field_set & breach_fields)
        union = len(field_set | breach_fields)
        field_similarity = overlap / union if union > 0 else 0

        # Sector match bonus
        sector_bonus = 0.3 if breach["sector"] == sector else 0

        # Scale similarity (closer user count = more relevant)
        if breach["records"] > 0 and num_users > 0:
            scale_ratio = min(breach["records"], num_users) / max(breach["records"], num_users)
        else:
            scale_ratio = 0

        total_score = field_similarity * 0.5 + sector_bonus + scale_ratio * 0.2

        if overlap > 0:  # Only include if at least 1 field overlaps
            scored.append({
                **breach,
                "similarity_score": round(total_score * 100, 1),
                "overlap_fields": sorted(list(field_set & breach_fields)),
                "overlap_count": overlap
            })

    scored.sort(key=lambda x: x["similarity_score"], reverse=True)
    return scored[:top_n]

# ============================================================
# FINANCIAL IMPACT CALCULATOR (from sector_benchmarks.csv)
# ============================================================

FIELD_COST_MULTIPLIER = {
    "medical_records": 2.5, "credit_card": 2.0, "biometric": 2.2,
    "ssn_aadhaar": 1.8, "bank_account": 1.9, "password": 1.3,
    "passport": 1.6, "email": 1.0, "phone": 1.1, "dob": 1.2,
    "address": 1.1, "username": 0.8, "ip_address": 0.7
}

def estimate_financial_impact(num_users, sector, fields):
    bench = SECTOR_BENCHMARKS.get(sector, {})
    base = bench.get("per_record_cost", 4.0)
    max_mult = max(FIELD_COST_MULTIPLIER.get(f, 1.0) for f in fields)
    estimated = num_users * base * max_mult

    return {
        "min": f"${estimated * 0.7:,.0f}",
        "max": f"${estimated * 1.3:,.0f}",
        "likely": f"${estimated:,.0f}"
    }

# ============================================================
# VISUALISATION ENGINES
# ============================================================

def render_gauge(score, severity):
    color_map = {
        "Low": "green", "Medium": "orange",
        "High": "red", "Critical": "darkred"
    }
    fig = go.Figure(go.Indicator(
        mode="gauge+number",
        value=score,
        domain={"x": [0, 1], "y": [0, 1]},
        title={"text": f"Risk Level: {severity}", "font": {"size": 20}},
        gauge={
            "axis": {"range": [0, 10], "tickwidth": 2},
            "bar": {"color": color_map.get(severity, "red")},
            "steps": [
                {"range": [0, 3], "color": "#d4edda"},
                {"range": [3, 6], "color": "#fff3cd"},
                {"range": [6, 8], "color": "#ffeeba"},
                {"range": [8, 10], "color": "#f8d7da"}
            ],
            "threshold": {
                "line": {"color": "black", "width": 4},
                "thickness": 0.75,
                "value": score
            }
        }
    ))
    fig.update_layout(height=300, margin=dict(t=60, b=20, l=40, r=40))
    return fig


def render_timeline(breach_date, jurisdiction):
    today = datetime.now()
    breach_dt = datetime.strptime(breach_date, "%Y-%m-%d")

    deadlines = [
        ("🔴 Breach Occurs", breach_dt),
        ("🇮🇳 DPDPA: Notify Board (Immediate)", breach_dt),
        ("🇪🇺 GDPR: Notify DPA (72 hrs)", breach_dt + timedelta(hours=72)),
        ("🇮🇳 DPDPA: Notify Principals (72 hrs)", breach_dt + timedelta(hours=72)),
        ("🇪🇺 GDPR: Full Investigation (30 days)", breach_dt + timedelta(days=30)),
    ]

    fig = go.Figure()
    for i, (label, date) in enumerate(deadlines):
        is_overdue = date < today
        status = "⚠️ OVERDUE" if is_overdue else "✅ Pending"
        color = "red" if is_overdue else "green"

        fig.add_trace(go.Scatter(
            x=[date], y=[i],
            mode="markers+text",
            text=[f"{label}<br>{date.strftime('%d %b %Y')}<br>{status}"],
            textposition="middle right",
            marker=dict(size=18, color=color, symbol="diamond"),
            textfont=dict(size=11),
            showlegend=False
        ))

    fig.update_layout(
        title="⏱️ Regulatory Compliance Deadline Timeline",
        xaxis_title="Date",
        yaxis=dict(visible=False),
        height=350,
        margin=dict(t=50, b=40, l=20, r=20),
        plot_bgcolor="rgba(0,0,0,0)"
    )
    return fig

# ============================================================
# BREACH SCENARIO LIBRARY
# ============================================================

SCENARIOS = {
    "🏥 Healthcare Breach": {
        "fields": ["email", "phone", "medical_records", "dob"],
        "sector": "healthcare",
        "users": 150000,
        "org": "MediCare Hospital",
        "jurisdiction": "India + EU",
        "breach_date": "2024-11-15"
    },
    "🏦 Banking Breach": {
        "fields": ["email", "password", "credit_card", "bank_account"],
        "sector": "finance",
        "users": 500000,
        "org": "NationalBank Ltd",
        "jurisdiction": "India",
        "breach_date": "2024-12-01"
    },
    "🛒 E-commerce Breach": {
        "fields": ["email", "phone", "address", "password"],
        "sector": "retail",
        "users": 1200000,
        "org": "ShopEasy India",
        "jurisdiction": "India + EU",
        "breach_date": "2025-01-10"
    },
    "🏛️ Government Breach": {
        "fields": ["ssn_aadhaar", "passport", "phone", "dob"],
        "sector": "government",
        "users": 80000,
        "org": "State Revenue Department",
        "jurisdiction": "India",
        "breach_date": "2025-02-20"
    }
}

# ============================================================
# LLM ENGINE (Groq — Llama 3.3 70B)
# Context-Augmented Compliance Analysis
# ============================================================

def run_llm_analysis(fields, sector, org_name, num_users,
                     jurisdiction, breach_date, risk_score,
                     severity, gdpr_articles, dpdpa_sections,
                     similar_breaches):

    fields_str = ", ".join([f.replace("_", " ").title() for f in fields])
    gdpr_str = ", ".join(gdpr_articles)
    dpdpa_str = ", ".join(dpdpa_sections)

    # Build HIBP context from loaded CSV data
    hibp_lines = []
    for field in fields:
        if field in HIBP_STATS:
            s = HIBP_STATS[field]
            hibp_lines.append(
                f"- {field.replace('_', ' ').title()}: appeared in {s['breach_count']} "
                f"known breaches, biggest was {s['biggest_breach']} "
                f"({s['biggest_records']:,} records)"
            )
    hibp_context = "\n".join(hibp_lines)

    # Build similar breach context from PRC data
    similar_lines = []
    for b in similar_breaches[:3]:
        similar_lines.append(
            f"- {b['name']} ({b['year']}): {b['records']:,} records, "
            f"sector: {b['sector']}, fields: {', '.join(b['overlap_fields'])}"
        )
    similar_context = "\n".join(similar_lines) if similar_lines else "No similar breaches found."

    prompt = f"""
You are a senior cybersecurity compliance analyst.

A data breach has occurred. Here are the facts:

Organisation: {org_name}
Sector: {sector}
Breach Date: {breach_date}
Affected Users: {num_users:,}
Jurisdiction: {jurisdiction}
Breached Data Fields: {fields_str}
Risk Score: {risk_score}/10 (Severity: {severity})
GDPR Articles Triggered: {gdpr_str}
DPDPA Sections Triggered: {dpdpa_str}

Historical Breach Intelligence (HaveIBeenPwned database):
{hibp_context}

Similar Past Breaches (Privacy Rights Clearinghouse database):
{similar_context}

GDPR Reference:
{GDPR_RULES}

DPDPA 2023 Reference:
{DPDPA_RULES}

Your task: Return a JSON object with EXACTLY these keys:

{{
    "executive_summary": "3-4 sentence summary of the breach impact, referencing similar past breaches",
    "immediate_actions": ["action1", "action2", "action3", "action4"],
    "gdpr_analysis": "2-3 sentences explaining GDPR obligations for this specific breach",
    "dpdpa_analysis": "2-3 sentences explaining DPDPA obligations for this specific breach",
    "risk_justification": "2-3 sentences explaining why this risk score was assigned, citing historical data",
    "disclosure_letter": "Full professional breach notification letter addressed to affected individuals"
}}

Return ONLY valid JSON. No markdown. No explanation outside JSON.
"""

    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        max_tokens=4096
    )
    text = response.choices[0].message.content.strip()

    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]
    text = text.strip()

    return json.loads(text)

# ============================================================
# DOCX GENERATOR
# ============================================================

def generate_docx(org_name, breach_date, num_users, fields,
                  risk_score, severity, gdpr_articles,
                  dpdpa_sections, disclosure_letter,
                  immediate_actions, gdpr_analysis, dpdpa_analysis,
                  financial_impact, sector, similar_breaches):

    doc = Document()

    title = doc.add_heading("DATA BREACH INCIDENT REPORT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %H:%M')}")
    doc.add_paragraph(f"Organisation: {org_name}")
    doc.add_paragraph(f"Sector: {sector.title()}")
    doc.add_paragraph(f"Breach Date: {breach_date}")
    doc.add_paragraph(f"Affected Users: {num_users:,}")
    doc.add_paragraph("")

    doc.add_heading("RISK ASSESSMENT", level=1)
    doc.add_paragraph(f"Risk Score: {risk_score}/10")
    doc.add_paragraph(f"Severity: {severity}")

    doc.add_heading("ESTIMATED FINANCIAL IMPACT", level=1)
    doc.add_paragraph(f"Minimum Estimate: {financial_impact['min']}")
    doc.add_paragraph(f"Most Likely: {financial_impact['likely']}")
    doc.add_paragraph(f"Maximum Estimate: {financial_impact['max']}")
    bench = SECTOR_BENCHMARKS.get(sector, {})
    doc.add_paragraph(f"Source: {bench.get('source', 'IBM Cost of Data Breach Report 2024')}")

    doc.add_heading("BREACHED DATA FIELDS", level=1)
    for field in fields:
        sens = PII_SENSITIVITY.get(field, "Medium")
        exp = FIELD_EXPLANATIONS.get(field, "Potential privacy risk")
        cat = PII_CONFIG.get(field, {}).get("category", "Other")
        doc.add_paragraph(
            f"• {field.replace('_', ' ').title()} [{cat}] — "
            f"Sensitivity: {sens} — {exp}",
            style="List Bullet"
        )

    doc.add_heading("HISTORICAL BREACH INTELLIGENCE", level=1)
    doc.add_paragraph("Source: HaveIBeenPwned API (hibp_breaches_full.json)")
    for field in fields:
        if field in HIBP_STATS:
            s = HIBP_STATS[field]
            doc.add_paragraph(
                f"• {field.replace('_', ' ').title()}: {s['breach_count']} known breaches, "
                f"biggest: {s['biggest_breach']} ({s['biggest_records']:,} records)",
                style="List Bullet"
            )

    doc.add_heading("SIMILAR PAST BREACHES", level=1)
    doc.add_paragraph("Source: Privacy Rights Clearinghouse (privacy_rights_clearinghouse.csv)")
    for b in similar_breaches:
        doc.add_paragraph(
            f"• {b['name']} ({b['org']}, {b['year']}): {b['records']:,} records — "
            f"Similarity: {b['similarity_score']}% — "
            f"Overlapping fields: {', '.join(b['overlap_fields'])}",
            style="List Bullet"
        )

    doc.add_heading("REGULATORY OBLIGATIONS", level=1)
    doc.add_heading("GDPR", level=2)
    doc.add_paragraph(gdpr_analysis)
    for article in gdpr_articles:
        doc.add_paragraph(f"• {article}", style="List Bullet")

    # Fine precedents from enforcement tracker
    fine_refs = {f: GDPR_FINES[f] for f in fields if f in GDPR_FINES}
    if fine_refs:
        doc.add_heading("GDPR Fine Precedents (enforcementtracker.com)", level=2)
        for field, ref in fine_refs.items():
            doc.add_paragraph(
                f"• {field.replace('_', ' ').title()}: {ref['example_org']} "
                f"fined {ref['fine']} ({ref['year']}) — {ref['articles']}",
                style="List Bullet"
            )

    doc.add_heading("DPDPA 2023", level=2)
    doc.add_paragraph(dpdpa_analysis)
    for section in dpdpa_sections:
        doc.add_paragraph(f"• {section}", style="List Bullet")

    doc.add_heading("IMMEDIATE ACTIONS REQUIRED", level=1)
    for action in immediate_actions:
        doc.add_paragraph(f"• {action}", style="List Bullet")

    doc.add_heading("BREACH NOTIFICATION LETTER", level=1)
    doc.add_paragraph(disclosure_letter)

    doc.add_heading("DATA SOURCES", level=1)
    doc.add_paragraph("• PII Classification: data/pii_field_config.json", style="List Bullet")
    doc.add_paragraph("• Compliance Rules: data/compliance_rules.json", style="List Bullet")
    doc.add_paragraph("• Breach History: data/hibp_breaches_full.json (HaveIBeenPwned API)", style="List Bullet")
    doc.add_paragraph("• Fine Precedents: data/gdpr_enforcement_tracker.csv", style="List Bullet")
    doc.add_paragraph("• Sector Benchmarks: data/sector_benchmarks.csv (IBM 2024)", style="List Bullet")
    doc.add_paragraph("• Similar Breaches: data/privacy_rights_clearinghouse.csv", style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ============================================================
# STREAMLIT UI
# ============================================================

st.set_page_config(
    page_title="Dark Web Breach Impact Analyser",
    page_icon="🔐",
    layout="wide"
)

st.title("🔐 Dark Web Data Breach Impact Analyser")
st.caption("Team A8 — KLE Technological University | CEVI")

# Show loaded datasets in an expander
with st.expander("📂 Loaded Datasets", expanded=False):
    st.markdown(f"""
    | Dataset | Source | Records |
    |---|---|---|
    | PII Field Config | `data/pii_field_config.json` | {len(PII_CONFIG)} fields |
    | Compliance Rules | `data/compliance_rules.json` | GDPR + DPDPA |
    | HIBP Breach Data | `data/hibp_breaches_full.json` | 998 real breaches |
    | GDPR Enforcement | `data/gdpr_enforcement_tracker.csv` | {len(GDPR_FINES)} precedents |
    | Sector Benchmarks | `data/sector_benchmarks.csv` | {len(SECTOR_BENCHMARKS)} sectors |
    | Historical Breaches | `data/privacy_rights_clearinghouse.csv` | {len(PRC_BREACHES)} breaches |
    """)

st.divider()

with st.sidebar:
    st.header("📋 Breach Details")

    st.subheader("⚡ Quick Scenarios")
    for scenario_name in SCENARIOS:
        if st.button(scenario_name, use_container_width=True):
            s = SCENARIOS[scenario_name]
            st.session_state["fields"] = s["fields"]
            st.session_state["sector"] = s["sector"]
            st.session_state["users"] = s["users"]
            st.session_state["org"] = s["org"]
            st.session_state["jurisdiction"] = s["jurisdiction"]
            st.session_state["breach_date"] = s["breach_date"]

    st.divider()
    st.subheader("✏️ Manual Input")

    org_name = st.text_input(
        "Organisation Name",
        value=st.session_state.get("org", "")
    )
    breach_date = st.text_input(
        "Breach Date (YYYY-MM-DD)",
        value=st.session_state.get("breach_date", "2025-01-01")
    )
    num_users = st.number_input(
        "Estimated Affected Users",
        min_value=1,
        value=st.session_state.get("users", 10000)
    )
    sector = st.selectbox(
        "Organisation Sector",
        ["healthcare", "finance", "government", "technology",
         "education", "retail"],
        index=["healthcare", "finance", "government", "technology",
               "education", "retail"].index(
            st.session_state.get("sector", "technology")
        )
    )
    jurisdiction = st.selectbox(
        "Jurisdiction",
        ["India", "EU", "India + EU"],
        index=["India", "EU", "India + EU"].index(
            st.session_state.get("jurisdiction", "India")
        )
    )

    st.subheader("🗂️ Breached Data Fields")
    all_fields = list(FIELD_WEIGHTS.keys())
    default_fields = st.session_state.get("fields", [])
    selected_fields = st.multiselect(
        "Select all fields found in the breach",
        options=all_fields,
        default=default_fields,
        format_func=lambda x: f"{x.replace('_', ' ').title()} ({PII_SENSITIVITY.get(x, 'Medium')})"
    )

    analyse_btn = st.button(
        "🔍 Run Impact Analysis",
        type="primary",
        use_container_width=True
    )

# ============================================================
# MAIN ANALYSIS FLOW
# ============================================================

if analyse_btn and selected_fields and org_name:
    logger.info(f"Streamlit: Commencing manual breach analysis for '{org_name}' [Sector: {sector}, Jurisdiction: {jurisdiction}, Users: {num_users:,}]")
    start_time = time.time()
    try:
        with st.spinner("Analysing breach... loading datasets & running AI compliance engine..."):

            risk_score, severity, color = calculate_risk(
                selected_fields, sector, num_users
            )
            gdpr_articles, dpdpa_sections = get_compliance_triggers(
                selected_fields
            )
            explanations = generate_explanation(selected_fields)
            financial_impact = estimate_financial_impact(
                num_users, sector, selected_fields
            )
            similar_breaches = find_similar_breaches(
                selected_fields, sector, num_users
            )

            logger.info("Streamlit: Submitting payload to Groq Llama 3.3 model...")
            llm_start_time = time.time()
            llm_output = run_llm_analysis(
                selected_fields, sector, org_name, num_users,
                jurisdiction, breach_date, risk_score, severity,
                gdpr_articles, dpdpa_sections, similar_breaches
            )
            logger.info(f"Streamlit: Groq LLM API response received in {(time.time() - llm_start_time)*1000:.2f}ms")

        duration_ms = (time.time() - start_time) * 1000
        # Record structured log in persistent log database
        log_incident(logger, org_name, sector, num_users, risk_score, severity, duration_ms)
        
    except Exception as e:
        logger.error(f"Streamlit: Critical failure during breach analysis execution: {str(e)}", exc_info=True)
        st.error(f"Breach Impact Engine Error: {str(e)}")
        llm_output = None

    if llm_output:
        # ========================================================
        # TABS
        # ========================================================

        tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs([
            "📊 Risk Score",
            "🗂️ PII Classification",
            "⚖️ Regulatory Obligations",
            "🔍 Similar Breaches",
            "📝 Disclosure Letter",
            "📄 Download Report",
            "📊 Historical Intelligence",
            "🕵️ Audit Log Viewer"
        ])

        # ------ TAB 1: RISK SCORE ------
        with tab1:
            st.subheader("Breach Risk Assessment")
            col1, col2, col3 = st.columns(3)
            col1.metric("Risk Score", f"{risk_score}/10")
            col2.metric("Severity", severity)
            col3.metric("Affected Users", f"{num_users:,}")

            # Gauge chart
            gauge = render_gauge(risk_score, severity)
            st.plotly_chart(gauge, use_container_width=True)

            st.divider()

            # Sector Benchmark (from sector_benchmarks.csv)
            st.subheader("📈 Sector Benchmark")
            bench = SECTOR_BENCHMARKS.get(sector, {})
            if bench:
                b1, b2, b3, b4 = st.columns(4)
                b1.metric("Industry Avg Breach Cost", f"${bench['avg_breach_cost_usd']}")
                b2.metric("Avg Detection Time", f"{bench['avg_detection_days']} days")
                b3.metric("Avg Containment", f"{bench['avg_containment_days']} days")
                b4.metric("Most Common Attack", bench["most_common_attack"])
                st.caption(f"Source: {bench['source']} | Data: data/sector_benchmarks.csv")

            st.divider()

            # Financial Impact (calculated from sector_benchmarks.csv per-record costs)
            st.subheader("💸 Estimated Financial Impact")
            fi1, fi2, fi3 = st.columns(3)
            fi1.metric("Minimum Estimate", financial_impact["min"])
            fi2.metric("Most Likely", financial_impact["likely"])
            fi3.metric("Maximum Estimate", financial_impact["max"])
            st.caption("Calculated using per-record costs from data/sector_benchmarks.csv")

            st.divider()

            # Executive Summary
            st.subheader("📋 Executive Summary")
            st.info(llm_output["executive_summary"])

            st.subheader("⚠️ Risk Justification")
            st.write(llm_output["risk_justification"])

            st.subheader("🚨 Immediate Actions Required")
            for action in llm_output["immediate_actions"]:
                st.error(f"→ {action}")

            st.divider()
            st.metric(
                "⚡ Time to Notify (DPDPA Rule 7)",
                "72 Hours",
                delta="Immediate board notification required",
                delta_color="inverse"
            )

        # ------ TAB 2: PII CLASSIFICATION ------
        with tab2:
            st.subheader("PII Sensitivity Classification")
            st.caption("Data source: data/pii_field_config.json + data/hibp_breaches_full.json")
            for item in explanations:
                sens = item["sensitivity"]
                icon = {"Critical": "🔴", "High": "🟠",
                        "Medium": "🟡", "Low": "🟢"}.get(sens, "⚪")
                raw_field = item["raw_field"]
                with st.expander(
                    f"{icon} {item['field']} [{item['category']}] — {sens} Sensitivity"
                ):
                    st.write(f"**Risk:** {item['risk']}")
                    # HIBP stats from loaded CSV
                    if raw_field in HIBP_STATS:
                        stats = HIBP_STATS[raw_field]
                        hc1, hc2, hc3 = st.columns(3)
                        hc1.metric("Known Breaches", f"{stats['breach_count']}")
                        hc2.metric("Biggest Breach", stats["biggest_breach"])
                        hc3.metric("Avg Records Lost", f"{stats['avg_records']:,}")
                        st.caption(f"Source: {stats['source']}")

        # ------ TAB 3: REGULATORY OBLIGATIONS ------
        with tab3:
            st.subheader("Regulatory Obligations")
            st.caption("Rules: data/compliance_rules.json | Fines: data/gdpr_enforcement_tracker.csv")

            # Compliance Timeline
            timeline = render_timeline(breach_date, jurisdiction)
            st.plotly_chart(timeline, use_container_width=True)

            st.divider()

            col1, col2 = st.columns(2)
            with col1:
                st.subheader("🇪🇺 GDPR")
                st.write(llm_output["gdpr_analysis"])
                st.markdown("**Articles Triggered:**")
                for article in gdpr_articles:
                    st.warning(article)

            with col2:
                st.subheader("🇮🇳 DPDPA 2023")
                st.write(llm_output["dpdpa_analysis"])
                st.markdown("**Sections Triggered:**")
                for section in dpdpa_sections:
                    st.error(section)

            # GDPR Fine Precedents from enforcement tracker CSV
            st.divider()
            st.subheader("💰 Real GDPR Fine Precedents")
            st.caption("Source: data/gdpr_enforcement_tracker.csv (enforcementtracker.com)")
            fine_refs = {f: GDPR_FINES[f] for f in selected_fields if f in GDPR_FINES}
            if fine_refs:
                for field, ref in fine_refs.items():
                    with st.expander(
                        f"⚖️ {field.replace('_', ' ').title()} → "
                        f"{ref['example_org']} fined {ref['fine']}"
                    ):
                        fc1, fc2, fc3 = st.columns(3)
                        fc1.write(f"**Year:** {ref['year']}")
                        fc2.write(f"**Authority:** {ref['authority']}")
                        fc3.write(f"**Country:** {ref['country']}")
                        st.write(f"**Articles Violated:** {ref['articles']}")
                        st.write(f"**Fine Amount:** {ref['fine']}")
            else:
                st.info("No direct fine precedents found for the selected data fields.")

        # ------ TAB 4: SIMILAR BREACHES (Privacy Rights Clearinghouse) ------
        with tab4:
            st.subheader("🔍 Similar Historical Breaches")
            st.caption("Source: data/privacy_rights_clearinghouse.csv | "
                       "Matching algorithm: Jaccard field similarity + sector + scale")

            if similar_breaches:
                for i, breach in enumerate(similar_breaches, 1):
                    with st.expander(
                        f"{'🔴' if breach['similarity_score'] > 60 else '🟠' if breach['similarity_score'] > 40 else '🟡'} "
                        f"#{i} — {breach['name']} ({breach['org']}, {breach['year']}) — "
                        f"{breach['similarity_score']}% similar"
                    ):
                        sc1, sc2, sc3, sc4 = st.columns(4)
                        sc1.metric("Records Lost", f"{breach['records']:,}")
                        sc2.metric("Sector", breach['sector'].title())
                        sc3.metric("Country", breach['country'])
                        sc4.metric("Overlap Fields", f"{breach['overlap_count']}")

                        st.write(f"**Overlapping Data Types:** "
                                 f"{', '.join([f.replace('_', ' ').title() for f in breach['overlap_fields']])}")
                        st.write(f"**All Exposed Fields:** "
                                 f"{', '.join([f.replace('_', ' ').title() for f in breach['fields']])}")
            else:
                st.info("No similar breaches found in the database.")

            st.divider()
            st.info(
                f"💡 Searched {len(PRC_BREACHES)} historical breaches from the "
                f"Privacy Rights Clearinghouse database. "
                f"Found {len(similar_breaches)} breaches with overlapping data fields."
            )

        # ------ TAB 5: DISCLOSURE LETTER ------
        with tab5:
            st.subheader("Draft Breach Notification Letter")
            st.text_area(
                "Disclosure Statement",
                value=llm_output["disclosure_letter"],
                height=400
            )

        # ------ TAB 6: DOWNLOAD REPORT ------
        with tab6:
            st.subheader("Download Full Incident Report")
            st.caption("Includes: Risk assessment, PII classification, regulatory analysis, "
                       "historical intelligence, similar breaches, financial impact, "
                       "fine precedents, and disclosure letter.")
            st.markdown("**Data sources included in report:**")
            st.markdown("""
            - `pii_field_config.json` — PII field weights & sensitivity
            - `compliance_rules.json` — GDPR/DPDPA legal text & triggers
            - `hibp_breaches_full.json` — HaveIBeenPwned breach intelligence
            - `gdpr_enforcement_tracker.csv` — Real GDPR fine precedents
            - `sector_benchmarks.csv` — IBM 2024 sector cost data
            - `privacy_rights_clearinghouse.csv` — Historical breach records
            """)
            
            logger.info(f"Streamlit: Compiling docx report for downloading...")
            docx_start = time.time()
            docx_buffer = generate_docx(
                org_name, breach_date, num_users,
                selected_fields, risk_score, severity,
                gdpr_articles, dpdpa_sections,
                llm_output["disclosure_letter"],
                llm_output["immediate_actions"],
                llm_output["gdpr_analysis"],
                llm_output["dpdpa_analysis"],
                financial_impact, sector, similar_breaches
            )
            logger.info(f"Streamlit: Docx compiled successfully in {(time.time() - docx_start)*1000:.2f}ms")
            
            st.download_button(
                label="📥 Download Breach Report (.docx)",
                data=docx_buffer,
                file_name=f"breach_report_{org_name.replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument"
                     ".wordprocessingml.document"
            )

        # ------ TAB 7: HISTORICAL INTELLIGENCE ------
        with tab7:
            st.subheader("📊 Historical Breach Intelligence")
            st.caption("Source: data/hibp_breaches_full.json (HaveIBeenPwned — 998 documented breaches)")

            for field in selected_fields:
                if field in HIBP_STATS:
                    s = HIBP_STATS[field]
                    c1, c2, c3, c4 = st.columns(4)
                    c1.metric(
                        f"{field.replace('_', ' ').title()}",
                        f"{s['breach_count']} breaches"
                    )
                    c2.metric("Biggest Known Breach", s["biggest_breach"])
                    c3.metric("Most Recent Year", s["most_recent"])
                    c4.metric("Avg Records Lost", f"{s['avg_records']:,}")
                    st.divider()

            st.info(
                "💡 Historical data shows your breached fields have appeared in "
                "hundreds of real-world incidents, confirming high attacker interest "
                "in this data type. Data dynamically loaded and aggregated from data/hibp_breaches_full.json"
            )

        # ------ TAB 8: AUDIT LOG VIEWER ------
        with tab8:
            st.subheader("🕵️ Live System Audit Trail")
            st.caption("Real-time persistent event execution audit log loaded from `logs/breach_analytics.log`.")
            
            from logger import LOG_FILE_PATH
            if os.path.exists(LOG_FILE_PATH):
                import pandas as pd
                try:
                    with open(LOG_FILE_PATH, "r", encoding="utf-8") as f:
                        log_lines = f.readlines()
                    
                    parsed_rows = []
                    for line in reversed(log_lines):
                        line = line.strip()
                        if not line:
                            continue
                        parts = line.split(" | ")
                        if len(parts) >= 4:
                            timestamp = parts[0].strip()
                            lvl = parts[1].strip()
                            module = parts[2].strip()
                            msg = " | ".join(parts[3:])
                            
                            meta = ""
                            if " | {" in msg:
                                msg, meta = msg.split(" | {", 1)
                                meta = "{" + meta
                                
                            parsed_rows.append({
                                "Timestamp": timestamp,
                                "Level": lvl,
                                "Source": module,
                                "Event Message": msg,
                                "Structured Metadata": meta
                            })
                    
                    df = pd.DataFrame(parsed_rows)
                    if not df.empty:
                        col_f1, col_f2 = st.columns([1, 3])
                        with col_f1:
                            filter_lvl = st.selectbox("Log Severity Level", ["ALL", "INFO", "WARNING", "ERROR", "CRITICAL"])
                        with col_f2:
                            search_term = st.text_input("Grep Search Log Message", "")
                            
                        if filter_lvl != "ALL":
                            df = df[df["Level"] == filter_lvl]
                        if search_term:
                            # Safely search columns without raising errors on nulls
                            df = df[
                                df["Event Message"].str.contains(search_term, case=False, na=False) | 
                                df["Structured Metadata"].str.contains(search_term, case=False, na=False)
                            ]
                            
                        def color_level(val):
                            bg = 'rgba(0,0,0,0)'
                            fg = '#ffffff'
                            if val in ['ERROR', 'CRITICAL']:
                                bg = '#ffcccc'
                                fg = '#7f0000'
                            elif val == 'WARNING':
                                bg = '#ffeebb'
                                fg = '#7f5500'
                            elif val == 'INFO':
                                bg = '#e2f0d9'
                                fg = '#2e5b1e'
                            return f'background-color: {bg}; color: {fg}; font-weight: bold;'
                        
                        styled_df = df.style.applymap(color_level, subset=['Level'])
                        st.dataframe(styled_df, use_container_width=True, height=400)
                        
                        if st.button("🔄 Refresh Logs"):
                            st.rerun()
                    else:
                        st.info("Log file is currently empty.")
                except Exception as e:
                    st.error(f"Error parsing system log file: {str(e)}")
            else:
                st.info("No active logs recorded in the session yet.")

elif analyse_btn and not selected_fields:
    st.warning("Please select at least one breached data field.")
elif analyse_btn and not org_name:
    st.warning("Please enter the organisation name.")
else:
    st.info("👈 Select a quick scenario or fill in breach details, then click Run Impact Analysis.")
